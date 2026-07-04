import spacy
#----------------------------
from cefrpy import CEFRSpaCyAnalyzer,CEFRLevel
#----------------------------
import pymupdf
#----------------------------
from nltk.corpus import wordnet
from nltk.corpus.reader.wordnet import NOUN, VERB, ADJ, ADV
#----------------------------
import re
import unicodedata
import json
import os
import sys
#----------------------------
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
'''
*Notes:
**
it doesnt get the words that are the same level as the book and it shows the harder words
for example if the CEFR level of the book is B1, it gets the B2 and higher level words

**
book.get_CEFR_level()[0] -> it shows the float of the CEFR level
book.get_CEFR_level()[1] -> shows the letter of the CEFR level (A1,A2,B1,B2,C1,C2)


to do list:
get pages of pdf file (done)

Troubleshoots:
__Troubleshooting CEFR levels of words__

'''
class Komaki:
    def __init__(self,file_name):
        self.file_full_name = os.path.basename(file_name)
        self.file_name_only = os.path.splitext(self.file_full_name)[0]
        self.dir = os.path.dirname(file_name)
        # try:
        #     nltk.data.find('corpora/wordnet')
        # except LookupError:
        #     nltk.download('wordnet')
        #loading spacy's model
        self.nlp = spacy.load("en_core_web_md")
        # self.nlp = self.load_spacy_model()
        self.spacy_analyzer = CEFRSpaCyAnalyzer()
        #opening the file:
        #   pdf file:
        if file_name.endswith(".pdf"):
            self.book_pages = self.get_pdf_pages(file_name)
        
        #____Configurations____
        self.footer_font = 7
        self.margin = 20
        self.scale_factor_1 = 7
        self.scale_factor_2 = 8
        self.text_font_size = 15
        # 6/10 of text font size 
        self.dict_font_size = 9
    
    #1
    # def load_spacy_model(self):
    #     try:
    #         # Try loading normally first
    #         return spacy.load("en_core_web_lg")
    #     except OSError:
    #         # If that fails, try loading from the PyInstaller bundle
    #         try:
    #             if hasattr(sys, '_MEIPASS'):
    #                 model_path = os.path.join(sys._MEIPASS, 'en_core_web_lg')
    #                 return spacy.load(model_path)
    #             else:
    #                 raise
    #         except:
    #             # Fallback to a smaller model if available
    #             try:
    #                 return spacy.load("en_core_web_sm")
    #             except:
    #                 raise OSError("No spaCy model found. Please install en_core_web_lg or en_core_web_sm")
    
    #2
    # def load_spacy_model(self):
    #     model_names = ["en_core_web_lg", "en_core_web_sm"]
        
    #     for model_name in model_names:
    #         try:
    #             if hasattr(sys, '_MEIPASS'):
    #                 # Running as PyInstaller bundle
    #                 model_path = os.path.join(sys._MEIPASS, model_name)
    #                 return spacy.load(model_path)
    #             else:
    #                 # Running normally
    #                 return spacy.load(model_name)
    #         except OSError:
    #             continue
        
    #     raise OSError("No spaCy model found. Please install en_core_web_lg or en_core_web_sm")

    #3
    # def load_spacy_model(self):
        
    #     print("=== SpaCy Model Debug Info ===")
        
    #     if hasattr(sys, '_MEIPASS'):
    #         print(f"Running in PyInstaller mode")
    #         print(f"_MEIPASS path: {sys._MEIPASS}")
            
    #         # List everything in _MEIPASS
    #         print("Contents of _MEIPASS:")
    #         try:
    #             for item in os.listdir(sys._MEIPASS):
    #                 print(f"  {item}")
    #                 if item.startswith('en_core_web'):
    #                     item_path = os.path.join(sys._MEIPASS, item)
    #                     if os.path.isdir(item_path):
    #                         print(f"    Contents of {item}:")
    #                         for subitem in os.listdir(item_path)[:10]:  # Show first 10 items
    #                             print(f"      {subitem}")
    #         except Exception as e:
    #             print(f"Error listing _MEIPASS: {e}")
    #     else:
    #         print("Running in normal mode")
        
    #     model_names = ["en_core_web_lg", "en_core_web_sm"]
        
    #     for model_name in model_names:
    #         print(f"\nTrying to load: {model_name}")
    #         try:
    #             if hasattr(sys, '_MEIPASS'):
    #                 model_path = os.path.join(sys._MEIPASS, model_name)
    #                 print(f"  Trying path: {model_path}")
    #                 print(f"  Path exists: {os.path.exists(model_path)}")
    #                 if os.path.exists(model_path):
    #                     print(f"  Path contents: {os.listdir(model_path)[:5]}")
    #                 return spacy.load(model_path)
    #             else:
    #                 print(f"  Loading normally: {model_name}")
    #                 return spacy.load(model_name)
    #         except Exception as e:
    #             print(f"  Failed: {e}")
    #             continue
        
    #     raise OSError("No spaCy model found. Please install en_core_web_lg or en_core_web_sm")

    # def load_spacy_model(self):

    #     model_names = ["en_core_web_lg", "en_core_web_sm"]
        
    #     for model_name in model_names:
    #         print(f"Trying to load: {model_name}")
            
    #         try:
    #             # First try: Load from system installation (most reliable)
    #             print(f"  Trying system installation...")
    #             return spacy.load(model_name)
    #         except OSError as e:
    #             print(f"  System load failed: {e}")
            
    #         if hasattr(sys, '_MEIPASS'):
    #             try:
    #                 # Second try: Load from PyInstaller bundle
    #                 model_path = os.path.join(sys._MEIPASS, model_name)
    #                 print(f"  Trying PyInstaller bundle: {model_path}")
    #                 if os.path.exists(model_path):
    #                     return spacy.load(model_path)
    #             except OSError as e:
    #                 print(f"  Bundle load failed: {e}")
            
    #         try:
    #             # Third try: Try to import as module and get path
    #             print(f"  Trying module import...")
    #             import importlib
    #             mod = importlib.import_module(model_name)
    #             return spacy.load(mod.__path__[0])
    #         except Exception as e:
    #             print(f"  Module import failed: {e}")
        
    #     # If all else fails, show clear instructions
    #     raise OSError(f"""
    #     No spaCy model found. Please install one by running:
        
    #     python -m spacy download en_core_web_lg
        
    #     Then run this program again. The model will be loaded from your system installation.
    #     """)


    def get_pdf_pages(self,file_name):
        # print("get_pdf_pages function is ok!")
        doc = pymupdf.open(file_name)
        return doc
    
    def get_CEFR_level(self):
        levels_sum = 0.0
        words_array = []
        filtered_tokens = []
        words_number = 0
        for page in self.book_pages:
            text = page.get_text()
            #____cleaning the text____
            text = self.clean_extracted_text(text)
            #____cleaning the text____
            doc = self.nlp(text)
            # analyzed tokens are words that are arrays themselves, its a array that has arrays in it
            analyzed_tokens = self.spacy_analyzer.analize_doc(doc)
            for token_data in analyzed_tokens:
                word,pos,is_skipped,level,start_pos,end_pos = token_data
                if level is None:
                    continue
                if pos == "PRON":
                    continue

                    # Get the actual SpaCy token using the position
                token = doc.char_span(start_pos, end_pos)
                if token is None:
                    continue
                token = token[0]

                # Filter out stop words, punctuations, digits
                if token.is_stop or token.is_punct or token.like_num:
                    continue

                # filtered_tokens.append(token_data)
                if word in words_array:
                    continue

                words_array.append(word)
                levels_sum += level
                words_number += 1
                # print(levels_sum , words_number)
        cefrlevel_float = round(levels_sum/words_number)
        cefrlevel = CEFRLevel(cefrlevel_float)
        # print(round(levels_sum/words_number))
        # print(cefrlevel)
        return [cefrlevel_float,cefrlevel]

    def cefr_to_wordnet_pos(self,tag):
        if tag in {'NN', 'NNS', 'NNP', 'NNPS'}:
            return NOUN
        elif tag in {'VB', 'VBD', 'VBG', 'VBN', 'VBP', 'VBZ'}:
            return VERB
        elif tag in {'JJ', 'JJR', 'JJS'}:
            return ADJ
        elif tag in {'RB', 'RBR', 'RBS', 'WRB'}:
            return ADV
        else:
            return None  # No WordNet equivalent

    
    def clean_extracted_text(self,text):
        if not text:
            return ""
        
        # Handle common PDF extraction issues
        # Remove null bytes and control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        
        # Normalize Unicode
        text = unicodedata.normalize('NFKD', text)
        
        # Replace characters outside the safe range (0-255)
        cleaned_chars = []
        for char in text:
            char_code = ord(char)
            if char_code <= 255:
                cleaned_chars.append(char)
            else:
                # Replace with space or closest ASCII equivalent
                cleaned_chars.append(' ')
        
        text = ''.join(cleaned_chars)
        
        # Clean up multiple spaces and normalize whitespace
        text = re.sub(r'\s+', ' ', text)
        
        return text.strip()
    #__________________________________get_difficult_words__________________________________
    def get_difficult_words(self,cefrlevel_float):
        # #__Troubleshooting CEFR levels of words__
        # ts = ""
        # #__Troubleshooting CEFR levels of words__
        pages_dic = []
        page_num = 0
        for page in self.book_pages:
            page_dic = {}
            analyzed_words = []
            # a dictionary of word definitions for every page
            text = page.get_text()
            #____cleaning the text____
            text = self.clean_extracted_text(text)
            #____cleaning the text____
            doc = self.nlp(text)
            # analyzed tokens are words that are arrays themselves, its a array that has arrays in it
            analyzed_tokens = self.spacy_analyzer.analize_doc(doc)
            #___________________
            # for sent in doc.sents:
            #     sent.spa
            #___________________
            for token_data in analyzed_tokens:
                word,pos,is_skipped,level,start_pos,end_pos = token_data
                if level is None:
                    continue
                if pos == "PRON":
                    continue

                    # Get the actual SpaCy token using the position
                token = doc.char_span(start_pos, end_pos)
                if token is None:
                    continue
                token = token[0]

                if token.lemma_.lower() in analyzed_words:
                    continue

                # Filter out stop words, punctuations, digits
                if token.is_stop or token.is_punct or token.like_num:
                    continue
                
                # #__Troubleshooting CEFR levels of words__
                # ts += f"{word}:{level}\n"
                # #__Troubleshooting CEFR levels of words__

                # if the word is what we are looking for
                if round(level) > cefrlevel_float:
                    wn_pos = self.cefr_to_wordnet_pos(pos)
                    if wn_pos:
                        synsets = wordnet.synsets(token.lemma_,wn_pos)
                        analyzed_words.append(token.lemma_.lower())
                        if synsets:
                            # get the best sysnsets based on synset and sentence similarity
                            for sent in doc.sents:
                                if sent.start_char <= start_pos < sent.end_char:
                                    sent_doc = self.nlp(sent.text)
                                    score = []
                                    for synset in synsets:
                                        #CHANGED
                                        word_sense = synset.definition()
                                        for example in synset.examples():
                                            word_sense += " "+ example
                                        synset_doc = self.nlp(word_sense)
                                        #CHANGED
                                        similarity = sent_doc.similarity(synset_doc)
                                        score.append((synset,similarity))
                                    best_syn = sorted(score,key=lambda x: x[1], reverse=True)[0][0]
                                    #__________________ get the best sysnonym for the best_syn
                                    synonyms = []
                                    similarity = 0
                                    for lemma in best_syn.lemmas():
                                        if lemma.name() != token.lemma_.lower():
                                            # print(lemma.name())
                                            synonym_doc = self.nlp(lemma.name())
                                            similarity = sent_doc.similarity(synonym_doc)
                                            synonyms.append((lemma.name(),similarity))
                                    # if there were no synonyms avaible, use its hypernyms
                                    if not synonyms:
                                        for hypernyms in best_syn.hypernyms():
                                            for hypernym in hypernyms.lemmas():
                                                hypernym_doc = self.nlp(hypernym.name())
                                                similarity = sent_doc.similarity(hypernym_doc)
                                                synonyms.append((hypernym.name(),similarity))
                                    # if the synonyms were still empty, give the synonyms array an empty space
                                    if not synonyms:
                                        synonyms.append(("",0))
                                    best_synonym = sorted(synonyms,key=lambda x: x[1], reverse=True)[0][0]
                                    # print(best_synonym)
                                    #__________________
                                    page_dic[token.text] = f"{best_synonym}/{best_syn.definition()}"
                                    # print(page_num,"\n\n\n_SENT_:",sent.text,"\n_WORD_:",token.text,"\n_SYN:",best_synonym,"\n_DEF_:",best_syn.definition())
            # print("\n\n",page_num)
            # print(page_dic)
            pages_dic.append(page_dic)
            page_num += 1
        # print("__________\n",pages_dic)

        # #__Troubleshooting CEFR levels of words__
        # with open("words_levels.txt","w+") as tsf:
        #     tsf.write(ts)
        # #__Troubleshooting CEFR levels of words__

        return pages_dic
    #__________________________________get_difficult_words__________________________________
    
    def make_processed_pdf(self,pages_dic):
        new_doc = pymupdf.open()
        for page_num,page in enumerate(self.book_pages):
            page_dic = pages_dic[page_num]
            footer = ""
            for key,value in page_dic.items():
                footer += f" | {key}: {value}"
            # print(footer)
            width = page.rect.width
            height = page.rect.height
            #margin = 40
            margin = self.margin
            scale_factor = self.scale_factor_1/self.scale_factor_2
            #______footer heights______
            #______footer heights______
            # footer_height = height * (1 - scale_factor)
            footer_height = 160
            textbox_rect = pymupdf.Rect(margin,height-footer_height+17,width-margin,height-2)
            dictionary_rect = pymupdf.Rect(0,height-footer_height,width,height)
            footer_rect = pymupdf.Rect(0,height-footer_height+15,width,height)
            #______
            original_rect = page.rect
            new_page = new_doc.new_page(width=original_rect.width, height=original_rect.height)
            scaled_width = original_rect.width * scale_factor #7/8 * width
            scaled_height = original_rect.height * scale_factor
            center_margin = (1-scale_factor)*original_rect.width/2
            target_rect = pymupdf.Rect(center_margin, 0, scaled_width+center_margin, scaled_height)
            new_page.show_pdf_page(target_rect, self.book_pages, page_num)
            new_page.draw_rect(target_rect,color=(1.0,1.0,1.0),width =0.5)
            #______
            new_page.draw_rect(dictionary_rect,color=(0.0,0.0,0.0),fill=(0.0,0.0,0.0),width=1.0)
            new_page.draw_rect(footer_rect,color=(0.0,0.0,0.0),fill=(0.9,0.9,0.9),width=1.0)
            new_page.draw_rect(textbox_rect,color=(0.9,0.9,0.9),fill=(0.9,0.9,0.9))
            # page.insert_text((x,y),footer,fontsize=9,fontname="helv")
            #__________________________________________________________
            for key, value in page_dic.items():
                text_instances = new_page.search_for(key)
                for inst in text_instances:
                    try:
                        # Convert quad to rectangle - more robust
                        rect = pymupdf.Rect(inst)
                        if rect.width > 0 and rect.height > 0:  # Valid area
                            highlight = new_page.add_highlight_annot(rect)
                            highlight.update()
                    except (ValueError, TypeError) as e:
                        print(f"Skipping invalid highlight for '{key}': {e}")
            #__________________________________________________________
            new_page.insert_textbox(dictionary_rect,"\tDictionary",fontsize=self.footer_font+1,fontname="helv",align=0,color=(1,1,1))
            new_page.insert_textbox(textbox_rect,footer,fontsize=self.footer_font,fontname="helv",align=0)
            for key,value in page_dic.items():
                text_instances = new_page.search_for(f"{key}:")
                for inst in text_instances:
                    try:
                        # Convert quad to rectangle - more robust
                        rect = pymupdf.Rect(inst)
                        if rect.width > 0 and rect.height > 0:  # Valid area
                            highlight = new_page.add_highlight_annot(rect)
                            highlight.update()
                    except (ValueError, TypeError) as e:
                        print(f"Skipping invalid highlight for '{key}': {e}")
        new_doc.save(f"{self.dir}/{self.file_name_only}_komaki.pdf")
        # self.book_pages.save("walden_with_footer.pdf")
        new_doc.close()
        self.book_pages.close()
    
    def make_line_by_line(self,pages_dic):

        #_____________________bolding words in paragraph_______________________
        def add_bold_words(paragraph, sentence, bold_words):
            words = sentence.split()
            for word in words:
                run = paragraph.add_run(word + " ")
                run.font.size = Pt(self.text_font_size)
                #Change#1
                if word.strip('.,!?\"\':;*()+=\\/<>}{][') in bold_words:
                    run.bold = True
        #_____________________bolding words in paragraph_______________________

        def set_paragraph_background(paragraph, color="D3D3D3"):
            """Set background shading for a paragraph (color is hex string, e.g., 'FFFF00')"""
            p = paragraph._p  # get the lxml element of the paragraph
            pPr = p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), color)
            pPr.append(shd)

        docx = Document()
        for page_num, page in enumerate(self.book_pages):
            page_dic = pages_dic[page_num]
            text = page.get_text()
            #_____FIXING_____
            text = self.clean_extracted_text(text)
            #_____FIXING_____
            doc = self.nlp(text)
            sents = ""
            for sent in doc.sents:
                sent_text = sent.text.replace("\n"," ")
                # docx.add_paragraph(sent_text)
                sents += sent_text+" "
                # print(sent.text)
                sent_dic = []
                for wrd in page_dic:
                    dic_wrd = re.findall(wrd,sent_text)
                    if dic_wrd:
                        sent_dic.append(dic_wrd[0])
                if sent_dic:
                    main_para = docx.add_paragraph()
                    add_bold_words(main_para,sents,sent_dic)
                    sents =""
                    for wrd in sent_dic:
                        para = docx.add_paragraph()
                        # making the difficult words bold
                        # run_word = para.add_run(f"{wrd}").bold = True
                        
                        run_word = para.add_run(f"{wrd}")
                        run_word.font.size = Pt(self.dict_font_size)
                        run_word.bold = True
                        
                        run_def = para.add_run(f" : {page_dic[wrd]}")
                        run_def.font.size = Pt(self.dict_font_size)

                        set_paragraph_background(para)
            # main_para = docx.add_paragraph(sents)
            main_para = docx.add_paragraph()
            run_main = main_para.add_run(sents)
            run_main.font.size = Pt(self.text_font_size)
        docx.save(f"{self.dir}/{self.file_name_only}_komaki.docx")
                
                    
                



                                    






if __name__ == "__main__":
    # file_name = "test5"
    file_name = input("Enter the address of the file:")
    book = Komaki(file_name)
    print(book.get_CEFR_level()[1],":",book.get_CEFR_level()[0])
    # print(book.get_difficult_words(book.get_CEFR_level()[0]))
    # print(book.get_difficult_words(4))
    print(book.file_name_only)
    if not os.path.exists(f"{book.file_name_only}.json"):
        json_data = book.get_difficult_words(4)
        with open(f"{book.file_name_only}.json","w") as file:
            json.dump(json_data,file,indent=2)
    with open(f"{book.file_name_only}.json","r") as f:
        dictionary = json.load(f)

    # book.make_processed_pdf(dictionary)
    book.make_line_by_line(dictionary)
